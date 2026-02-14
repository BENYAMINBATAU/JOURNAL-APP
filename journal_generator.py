"""
Journal Generator Module
Generates journal articles in DOCX or PDF format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Any
import re


class JournalGenerator:
    """Generate journal articles from thesis content"""
    
    def __init__(self, template: str = 'unm'):
        """
        Initialize generator
        
        Args:
            template: Template name (currently supports 'unm')
        """
        self.template = template
        self.page_width = Inches(8.27)  # A4
        self.page_height = Inches(11.69)  # A4
        self.margin = Inches(0.98)  # 2.5cm
        self.margin_top = Inches(1.18)  # 3cm
        
    def create_journal_article(self, thesis_content: Dict[str, Any], settings: Dict) -> Dict[str, Any]:
        """
        Create journal article structure from thesis content
        
        Args:
            thesis_content: Extracted and enhanced thesis content
            settings: User settings
            
        Returns:
            Journal article data structure
        """
        journal_data = {
            'title': self._create_title(thesis_content, settings),
            'authors': self._format_authors(settings),
            'affiliation': settings.get('affiliation', ''),
            'email': settings.get('email', ''),
            'abstract_en': thesis_content.get('abstract_en', ''),
            'keywords_en': thesis_content.get('keywords_en', ''),
            'abstract_id': thesis_content.get('abstract_id', ''),
            'keywords_id': thesis_content.get('keywords_id', ''),
            'introduction': self._create_introduction(thesis_content),
            'methodology': thesis_content.get('methodology_summary', thesis_content.get('methodology', '')),
            'results': thesis_content.get('results_summary', thesis_content.get('results', '')),
            'discussion': self._extract_discussion(thesis_content),
            'conclusions': thesis_content.get('conclusions_summary', thesis_content.get('conclusions', '')),
            'acknowledgments': self._create_acknowledgments(settings),
            'references': thesis_content.get('references', [])
        }
        
        return journal_data
    
    def _create_title(self, thesis_content: Dict, settings: Dict) -> str:
        """Create journal title from thesis title"""
        title = thesis_content.get('title', '').upper()
        
        # Remove common thesis words
        title = re.sub(r'TESIS|SKRIPSI|DISERTASI|THESIS|DISSERTATION', '', title, flags=re.IGNORECASE)
        
        # Clean up
        title = re.sub(r'\s+', ' ', title).strip()
        
        return title
    
    def _format_authors(self, settings: Dict) -> str:
        """Format author names"""
        author_name = settings.get('author_name', '')
        coauthors = settings.get('coauthors', '')
        
        if coauthors:
            return f"{author_name}¹*, {coauthors}"
        
        return f"{author_name}¹*"
    
    def _create_introduction(self, thesis_content: Dict) -> str:
        """Create introduction section"""
        background = thesis_content.get('background', '')
        objectives = thesis_content.get('objectives', '')
        
        # Combine and summarize
        intro = background[:1500]  # Limit length
        
        if objectives:
            intro += "\n\n" + objectives[:500]
        
        return intro
    
    def _extract_discussion(self, thesis_content: Dict) -> str:
        """Extract discussion from results"""
        results = thesis_content.get('results', '')
        
        # Try to find discussion section
        match = re.search(r'pembahasan(.*?)(?:simpulan|kesimpulan|$)', results, re.IGNORECASE | re.DOTALL)
        
        if match:
            return match.group(1).strip()[:2000]
        
        return results[:2000]
    
    def _create_acknowledgments(self, settings: Dict) -> str:
        """Create acknowledgments section"""
        affiliation = settings.get('affiliation', 'Universitas Negeri Makassar')
        
        return f"Peneliti mengucapkan terima kasih kepada {affiliation} serta semua pihak yang telah mendukung penelitian ini."
    
    def generate_docx(self, journal_data: Dict[str, Any], output_path: str):
        """
        Generate DOCX journal article
        
        Args:
            journal_data: Journal article data
            output_path: Path to save DOCX file
        """
        doc = Document()
        
        # Setup document
        self._setup_document(doc)
        
        # Add title
        self._add_title(doc, journal_data['title'])
        
        # Add authors
        self._add_authors(doc, journal_data['authors'], journal_data['affiliation'], journal_data['email'])
        
        # Add abstract (English)
        if journal_data.get('abstract_en'):
            self._add_section(doc, 'ABSTRACT', journal_data['abstract_en'], italic=True, center_title=True)
            if journal_data.get('keywords_en'):
                self._add_keywords(doc, journal_data['keywords_en'], italic=True)
        
        # Add abstrak (Indonesian)
        if journal_data.get('abstract_id'):
            self._add_section(doc, 'ABSTRAK', journal_data['abstract_id'], center_title=True)
            if journal_data.get('keywords_id'):
                self._add_keywords(doc, journal_data['keywords_id'])
        
        # Add introduction
        self._add_section(doc, 'LATAR BELAKANG', journal_data.get('introduction', ''))
        
        # Add methodology
        if journal_data.get('methodology'):
            self._add_section(doc, 'METODE PENELITIAN', journal_data['methodology'])
        
        # Add results and discussion
        if journal_data.get('results'):
            self._add_section(doc, 'HASIL DAN PEMBAHASAN', journal_data['results'])
        
        # Add conclusions
        if journal_data.get('conclusions'):
            self._add_section(doc, 'KESIMPULAN', journal_data['conclusions'])
        
        # Add acknowledgments
        if journal_data.get('acknowledgments'):
            self._add_section(doc, 'UCAPAN TERIMA KASIH', journal_data['acknowledgments'])
        
        # Add references
        if journal_data.get('references'):
            self._add_references(doc, journal_data['references'])
        
        # Save document
        doc.save(output_path)
    
    def _setup_document(self, doc: Document):
        """Setup document formatting"""
        # Set margins
        sections = doc.sections
        for section in sections:
            section.page_width = self.page_width
            section.page_height = self.page_height
            section.top_margin = self.margin_top
            section.bottom_margin = self.margin
            section.left_margin = self.margin
            section.right_margin = self.margin
        
        # Define styles
        styles = doc.styles
        
        # Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(11)
        
    def _add_title(self, doc: Document, title: str):
        """Add title"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.all_caps = True
        p.paragraph_format.space_after = Pt(12)
    
    def _add_authors(self, doc: Document, authors: str, affiliation: str, email: str):
        """Add authors and affiliation"""
        # Authors
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(authors)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        p.paragraph_format.space_after = Pt(6)
        
        # Affiliation
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"¹{affiliation}, {email}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(12)
    
    def _add_section(self, doc: Document, heading: str, content: str, italic: bool = False, center_title: bool = False):
        """Add a section with heading and content"""
        # Heading
        p = doc.add_paragraph()
        if center_title:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(heading)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        if italic:
            run.font.italic = True
        p.paragraph_format.space_after = Pt(6)
        
        # Content
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(content)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if italic:
            run.font.italic = True
        p.paragraph_format.space_after = Pt(12)
    
    def _add_keywords(self, doc: Document, keywords: str, italic: bool = False):
        """Add keywords"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Keywords label
        run = p.add_run('Keywords: ' if italic else 'Kata Kunci: ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        if italic:
            run.font.italic = True
        
        # Keywords text
        run = p.add_run(keywords)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if italic:
            run.font.italic = True
        
        p.paragraph_format.space_after = Pt(12)
    
    def _add_references(self, doc: Document, references: list):
        """Add references section"""
        # Heading
        p = doc.add_paragraph()
        run = p.add_run('DAFTAR PUSTAKA')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        p.paragraph_format.space_after = Pt(6)
        
        # References
        for ref in references[:25]:  # Limit to 25 references
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(ref)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0
    
    def generate_pdf(self, journal_data: Dict[str, Any], output_path: str):
        """
        Generate PDF journal article
        
        Args:
            journal_data: Journal article data
            output_path: Path to save PDF file
        """
        # First generate DOCX
        docx_path = output_path.replace('.pdf', '.docx')
        self.generate_docx(journal_data, docx_path)
        
        # Convert to PDF using reportlab or other library
        # For now, just keep as DOCX
        # TODO: Implement PDF conversion
        pass
